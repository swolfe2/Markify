
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    # Create the relationship
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a run for the text
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add blue color and underline for hyperlink styling
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_sample_docx(output_path):
    doc = Document()
    
    # Set narrow margins (0.5 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Title
    title = doc.add_heading('Markify Demo Document', 0)
    title.paragraph_format.space_after = Pt(4)
    
    doc.add_paragraph("This document showcases Markify's conversion capabilities.")

    # 1. Headers
    doc.add_heading('1. Header Support', level=1)
    doc.add_heading('1.1 Sub-header (Level 2)', level=2)
    doc.add_paragraph("Markify detects Word heading styles and emoji headers like 🔐 Security.")
    
    # 2. Text Formatting & Hyperlinks
    doc.add_heading('2. Formatting & Links', level=1)
    p = doc.add_paragraph()
    p.add_run("Text can be ")
    bold_run = p.add_run("bold")
    bold_run.bold = True
    p.add_run(" or ")
    italic_run = p.add_run("italic")
    italic_run.italic = True
    p.add_run(". Links work too: ")
    add_hyperlink(p, "Visit GitHub", "https://github.com")
    
    # 3. Lists
    doc.add_heading('3. Lists', level=1)
    doc.add_paragraph('Item A', style='List Bullet')
    doc.add_paragraph('Item B (nested below)', style='List Bullet')
    doc.add_paragraph('Sub-item B1', style='List Bullet 2')
    
    # 4. Tables
    doc.add_heading('4. Tables', level=1)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'ID', 'Name', 'Role'
    r1 = table.rows[1].cells
    r1[0].text, r1[1].text, r1[2].text = '001', 'Alice', 'Dev'
    r2 = table.rows[2].cells
    r2[0].text, r2[1].text, r2[2].text = '002', 'Bob', 'PM'
    
    # 5. Code Blocks - Power Query (M)
    doc.add_heading('5. Power Query (M)', level=1)
    code_table = doc.add_table(rows=1, cols=1)
    code_table.rows[0].cells[0].text = 'let\n    Source = Excel.Workbook(File.Contents("data.xlsx"))\nin\n    Source'
    
    # 6. Code Blocks - DAX
    doc.add_heading('6. DAX', level=1)
    dax_table = doc.add_table(rows=1, cols=1)
    dax_table.rows[0].cells[0].text = 'Revenue := SUMX(Sales, Sales[Qty] * Sales[Price])'
    
    # 7. Code Blocks - Python
    doc.add_heading('7. Python', level=1)
    py_table = doc.add_table(rows=1, cols=1)
    py_table.rows[0].cells[0].text = 'import pandas as pd\ndef load_data(path):\n    return pd.read_csv(path)'
    
    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Sample generated at: {output_path}")

if __name__ == '__main__':
    create_sample_docx("Samples/Markify_Demo.docx")
